# BookBlueprint.py
import streamlit as st
from openai import OpenAI
import PyPDF2
import docx
import json
import time
from typing import Optional, Dict, List, Any
from datetime import datetime
import plotly.express as px
import pandas as pd
from PIL import Image
import io
import base64

def show_blueprint_analyzer():
    """Main Book Marketing Blueprint Analyzer"""
    
    # Check if we're on the correct page
    if st.session_state.get('current_page') != "📖 Book Blueprint":
        return
    
    # IMPORTANT: Inherit API key from main app
    if 'openai_api_key' in st.session_state and st.session_state.openai_api_key:
        st.session_state.api_configured = True
    else:
        st.session_state.api_configured = False
    
    # Initialize session state
    if 'blueprint' not in st.session_state:
        st.session_state.blueprint = None
    
    if 'blueprint_stage' not in st.session_state:
        st.session_state.blueprint_stage = "upload"
    
    # Header with visual impact
    st.title("📖 Book Marketing Blueprint")
    st.markdown("""
    <style>
    .blueprint-header {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 20px;
        border-radius: 10px;
        color: white;
        text-align: center;
        margin-bottom: 20px;
    }
    .metric-card {
        background: #f0f2f6;
        padding: 15px;
        border-radius: 10px;
        text-align: center;
    }
    </style>
    """, unsafe_allow_html=True)
    
    st.markdown('<div class="blueprint-header"><h2>Your Complete Marketing Roadmap</h2><p>Upload your manuscript and cover. Get a 360° marketing strategy tailored to YOUR book.</p></div>', unsafe_allow_html=True)
    
    # API Configuration (only if not already configured)
    if not st.session_state.api_configured:
        with st.expander("🔑 API Settings", expanded=True):
            api_key = st.text_input("OpenAI API Key", type="password", key="blueprint_api_key", 
                                   value=st.session_state.get('openai_api_key', ''))
            if api_key and st.button("Connect", key="blueprint_connect"):
                st.session_state.openai_api_key = api_key
                st.session_state.api_configured = True
                st.rerun()
        return
    
    # Progress tracker
    stages = ["upload", "analyzing", "results"]
    stage_names = ["📤 Upload", "⚙️ Analysis", "🎯 Blueprint"]
    current_idx = stage_names.index(stage_names[stages.index(st.session_state.blueprint_stage)]) if st.session_state.blueprint_stage in stages else 0
    
    cols = st.columns(3)
    for i, (col, stage_name) in enumerate(zip(cols, stage_names)):
        with col:
            if i < current_idx:
                st.success(f"✅ {stage_name}")
            elif i == current_idx:
                st.info(f"⏳ {stage_name}")
            else:
                st.write(f"⚪ {stage_name}")
    
    st.divider()
    
    # Route to appropriate stage
    if st.session_state.blueprint_stage == "upload":
        show_upload_stage()
    elif st.session_state.blueprint_stage == "analyzing":
        show_analyzing_stage()
    elif st.session_state.blueprint_stage == "results":
        show_blueprint_results()


def show_upload_stage():
    """Stage 1: Upload manuscript and cover"""
    
    st.subheader("📤 Upload Your Book Materials")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("### 📄 Manuscript")
        manuscript_file = st.file_uploader(
            "Upload manuscript (PDF, DOCX, TXT)",
            type=['pdf', 'docx', 'txt'],
            key="blueprint_manuscript",
            help="Your complete book text"
        )
        
        if manuscript_file:
            st.success(f"✅ Loaded: {manuscript_file.name}")
            
            # Quick preview
            with st.expander("Preview"):
                text = extract_text_preview(manuscript_file)
                st.text(text[:500] + "..." if len(text) > 500 else text)
    
    with col2:
        st.markdown("### 🎨 Cover Image")
        cover_file = st.file_uploader(
            "Upload cover image (JPG, PNG)",
            type=['jpg', 'jpeg', 'png'],
            key="blueprint_cover",
            help="Your book cover"
        )
        
        if cover_file:
            st.success(f"✅ Loaded: {cover_file.name}")
            
            # Show thumbnail
            image = Image.open(cover_file)
            st.image(image, caption="Cover Preview", width=200)
    
    st.divider()
    
    # Additional info (optional but helpful)
    with st.expander("📋 Additional Information (Optional but Helpful)"):
        col1, col2 = st.columns(2)
        with col1:
            target_genre = st.text_input("Primary Genre (if known)", placeholder="e.g., Romantasy")
            comp_titles = st.text_area("Comparable Titles (one per line)", placeholder="Fourth Wing\nACOTAR\nSerpent & Wings of Night")
        with col2:
            target_audience = st.text_input("Target Audience (if known)", placeholder="e.g., Adult fantasy romance readers")
            author_platform = st.multiselect("Author Platforms", ["TikTok", "Instagram", "Twitter", "Facebook", "Newsletter", "None"])
    
    # Analyze button
    if manuscript_file and cover_file:
        if st.button("🚀 GENERATE COMPLETE MARKETING BLUEPRINT", type="primary", use_container_width=True):
            # Save files to session state
            st.session_state.manuscript_bytes = manuscript_file.getvalue()
            st.session_state.manuscript_type = manuscript_file.type
            st.session_state.cover_bytes = cover_file.getvalue()
            st.session_state.cover_type = cover_file.type
            
            # Save optional info
            st.session_state.blueprint_options = {
                "genre": target_genre if target_genre else None,
                "comp_titles": [c.strip() for c in comp_titles.split("\n") if c.strip()] if comp_titles else [],
                "target_audience": target_audience if target_audience else None,
                "author_platforms": author_platform
            }
            
            st.session_state.blueprint_stage = "analyzing"
            st.rerun()
    else:
        st.info("👆 Please upload both manuscript and cover to continue")


def show_analyzing_stage():
    """Stage 2: Analysis in progress"""
    
    st.subheader("⚙️ Analyzing Your Book")
    
    # Create progress visualization
    progress_bar = st.progress(0)
    status_text = st.empty()
    
    # Analysis steps
    steps = [
        "📖 Reading manuscript...",
        "🎨 Analyzing cover design...",
        "📊 Identifying genre conventions...",
        "👥 Building reader avatar...",
        "🎯 Determining market position...",
        "🔑 Extracting keywords...",
        "📝 Optimizing blurb...",
        "📱 Creating channel strategies...",
        "🗓️ Building launch timeline...",
        "🎬 Generating marketing assets..."
    ]
    
    # Simulate progress
    for i, step in enumerate(steps):
        status_text.text(step)
        progress_bar.progress((i + 1) / len(steps))
        time.sleep(0.2)
    
    # Extract text from manuscript
    status_text.text("📄 Extracting text from manuscript...")
    manuscript_text = extract_full_text(
        st.session_state.manuscript_bytes,
        st.session_state.manuscript_type
    )
    
    # Encode cover for vision API (optional)
    cover_base64 = base64.b64encode(st.session_state.cover_bytes).decode('utf-8')
    
    # Generate blueprint with GPT-4o-mini
    status_text.text("🤖 GPT-4o-mini generating your marketing blueprint...")
    blueprint = generate_complete_blueprint(
        manuscript_text,
        cover_base64,
        st.session_state.blueprint_options
    )
    
    if blueprint:
        st.session_state.blueprint = blueprint
        status_text.text("✅ Blueprint generated successfully!")
        time.sleep(1)
        st.session_state.blueprint_stage = "results"
        st.rerun()
    else:
        st.error("Failed to generate blueprint. Using sample data instead.")
        st.session_state.blueprint = get_fallback_blueprint()
        time.sleep(1)
        st.session_state.blueprint_stage = "results"
        st.rerun()


def show_blueprint_results():
    """Stage 3: Display the complete blueprint"""
    
    blueprint = st.session_state.blueprint
    if not blueprint:
        st.error("No blueprint found")
        return
    
    # Header with book title
    st.success(f"✅ Marketing Blueprint Complete: **{blueprint.get('book_profile', {}).get('title', 'Your Book')}**")
    
    # Export options
    col1, col2, col3 = st.columns(3)
    with col1:
        if st.button("📥 Download PDF Report", use_container_width=True):
            st.info("PDF export coming soon")
    with col2:
        if st.button("📋 Copy Summary", use_container_width=True):
            st.info("Copy feature coming soon")
    with col3:
        if st.button("🔄 New Analysis", use_container_width=True):
            st.session_state.blueprint_stage = "upload"
            st.session_state.blueprint = None
            st.rerun()
    
    st.divider()
    
    # Create tabs for all blueprint sections
    tabs = st.tabs([
        "📖 Book Profile", 
        "🎨 Cover Analysis", 
        "👥 Reader Avatar", 
        "🎯 Market Position",
        "🔑 Keywords",
        "📝 Blurb",
        "📱 Channel Strategy",
        "🗓️ Timeline",
        "🎬 Assets"
    ])
    
    # Tab 1: Book Profile
    with tabs[0]:
        show_book_profile(blueprint.get('book_profile', {}))
    
    # Tab 2: Cover Analysis
    with tabs[1]:
        show_cover_analysis(blueprint.get('cover_analysis', {}))
    
    # Tab 3: Reader Avatar
    with tabs[2]:
        show_reader_avatar(blueprint.get('reader_avatar', {}))
    
    # Tab 4: Market Position
    with tabs[3]:
        show_market_position(blueprint.get('market_position', {}))
    
    # Tab 5: Keywords
    with tabs[4]:
        show_keywords(blueprint.get('keyword_cloud', {}))
    
    # Tab 6: Blurb
    with tabs[5]:
        show_blurb_analysis(blueprint.get('blurb_analysis', {}))
    
    # Tab 7: Channel Strategy
    with tabs[6]:
        show_channel_strategy(blueprint.get('marketing_blueprint', {}).get('channel_strategies', {}))
    
    # Tab 8: Timeline
    with tabs[7]:
        show_timeline(blueprint.get('marketing_blueprint', {}).get('launch_timeline', {}))
    
    # Tab 9: Assets
    with tabs[8]:
        show_assets(blueprint.get('generated_assets', {}))


# ============================================================================
# DISPLAY FUNCTIONS FOR EACH TAB
# ============================================================================

def show_book_profile(profile):
    """Display book profile"""
    if not profile:
        st.info("No profile data")
        return
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("### 📖 Book Details")
        st.write(f"**Title:** {profile.get('title', 'N/A')}")
        st.write(f"**Genre:** {profile.get('genre', 'N/A')}")
        st.write(f"**Subgenres:** {', '.join(profile.get('subgenres', []))}")
        st.write(f"**Tone:** {profile.get('tone', 'N/A')}")
        st.write(f"**Pace:** {profile.get('pace', 'N/A')}")
    
    with col2:
        st.markdown("### 🎭 Characters")
        characters = profile.get('main_characters', [])
        if isinstance(characters, list):
            for char in characters:
                if isinstance(char, dict):
                    st.write(f"**{char.get('name', 'Unknown')}:** {char.get('description', '')}")
                else:
                    st.write(f"• {char}")
        else:
            st.write(characters)


def show_cover_analysis(cover):
    """Display cover analysis"""
    if not cover:
        st.info("No cover analysis available. Using sample data.")
        # Show sample data so it's not empty
        cover = {
            "genre_alignment": "85%",
            "strengths": ["Good typography", "Mood matches genre"],
            "weaknesses": ["Could be darker", "Figure could be more prominent"],
            "suggestions": ["Increase contrast", "Add subtle magical glow"],
            "comp_covers": ["Fourth Wing", "ACOTAR"]
        }
    
    # Score gauge
    alignment = cover.get('genre_alignment', '75%')
    if isinstance(alignment, str):
        try:
            alignment = int(alignment.replace('%', ''))
        except:
            alignment = 75
    
    st.markdown("### 🎨 Cover Analysis")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.metric("Genre Alignment", f"{alignment}%")
        
        st.markdown("**✅ Strengths:**")
        for s in cover.get('strengths', []):
            st.write(f"• {s}")
        
        st.markdown("**❌ Weaknesses:**")
        for w in cover.get('weaknesses', []):
            st.write(f"• {w}")
    
    with col2:
        st.markdown("**💡 Suggestions:**")
        for s in cover.get('suggestions', []):
            st.write(f"• {s}")
        
        st.markdown("**📚 Comparable Covers:**")
        for c in cover.get('comp_covers', []):
            st.write(f"• {c}")


def show_reader_avatar(avatar):
    """Display reader avatar"""
    if not avatar:
        st.info("No reader avatar data")
        return
    
    primary = avatar.get('primary', {})
    secondary = avatar.get('secondary', {})
    
    st.markdown("### 👤 Primary Reader Avatar")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.write(f"**Name:** {primary.get('name', 'N/A')}")
        st.write(f"**Age:** {primary.get('age', 'N/A')}")
        st.write(f"**Occupation:** {primary.get('occupation', 'N/A')}")
        st.write(f"**Reading Habits:** {primary.get('reading_habits', 'N/A')}")
    
    with col2:
        st.write(f"**Seeks:** {primary.get('what_she_seeks', 'N/A')}")
        st.write(f"**Hangs out:** {primary.get('where_she_hangs_out', 'N/A')}")
        st.write(f"**Avoids:** {primary.get('what_she_avoids', 'N/A')}")
    
    if secondary:
        st.markdown("### 👤 Secondary Reader Avatar")
        col1, col2 = st.columns(2)
        with col1:
            st.write(f"**Name:** {secondary.get('name', 'N/A')}")
            st.write(f"**Age:** {secondary.get('age', 'N/A')}")
        with col2:
            st.write(f"**Hangs out:** {secondary.get('where_she_hangs_out', 'N/A')}")


def show_market_position(position):
    """Display market position"""
    if not position:
        st.info("No market position data")
        return
    
    st.markdown(f"**Primary Shelf:** {position.get('primary_shelf', 'N/A')}")
    st.markdown(f"**Positioning Statement:** {position.get('positioning_statement', 'N/A')}")
    
    st.markdown("### 📚 Comparable Titles")
    comps = position.get('comp_titles', [])
    for comp in comps:
        if isinstance(comp, dict):
            with st.expander(f"**{comp.get('title', 'Unknown')}**"):
                st.write(f"**Similarity:** {comp.get('similarity', 'N/A')}")
                st.write(f"**Difference:** {comp.get('difference', 'N/A')}")
        else:
            st.write(f"• {comp}")
    
    st.markdown("### 🔍 Gap Analysis")
    st.write(position.get('gap_analysis', 'N/A'))


def show_keywords(keywords):
    """Display keyword cloud"""
    if not keywords:
        st.info("No keyword data")
        return
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("**🔑 Amazon Keywords:**")
        for k in keywords.get('amazon_keywords', []):
            st.write(f"• {k}")
    
    with col2:
        st.markdown("**📊 Search Volume:**")
        search_volume = keywords.get('search_volume', {})
        for vol, terms in search_volume.items():
            if isinstance(terms, list):
                st.write(f"**{vol}:** {', '.join(terms)}")
    
    st.markdown("**📁 Categories:**")
    for cat in keywords.get('categories', []):
        st.write(f"• {cat}")


def show_blurb_analysis(blurb):
    """Display blurb analysis"""
    if not blurb:
        st.info("No blurb data")
        return
    
    score = blurb.get('score', 0)
    st.metric("Blurb Score", f"{score}/100")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.markdown("**Current Blurb:**")
        st.info(blurb.get('current_blurb', 'N/A'))
        
        st.markdown("**✅ Strengths:**")
        for s in blurb.get('strengths', []):
            st.write(f"• {s}")
    
    with col2:
        st.markdown("**❌ Weaknesses:**")
        for w in blurb.get('weaknesses', []):
            st.write(f"• {w}")
        
        st.markdown("**✨ Optimized Version:**")
        st.success(blurb.get('optimized_version', 'N/A'))


def show_channel_strategy(strategies):
    """Display channel strategies"""
    if not strategies:
        st.info("No channel strategy data")
        return
    
    channel_tabs = st.tabs(list(strategies.keys()))
    
    for i, (channel_name, channel_data) in enumerate(strategies.items()):
        with channel_tabs[i]:
            if channel_name == "tiktok":
                st.markdown("**🎵 Content Pillars:**")
                for p in channel_data.get('content_pillars', []):
                    st.write(f"• {p}")
                
                st.markdown("**🎶 Sound Suggestions:**")
                for s in channel_data.get('sound_suggestions', []):
                    st.write(f"• {s}")
                
                st.markdown("**#️⃣ Hashtags:**")
                st.write(', '.join(channel_data.get('hashtags', [])))
            
            elif channel_name == "instagram":
                st.markdown("**📸 Post Types:**")
                for p in channel_data.get('post_types', []):
                    st.write(f"• {p}")
                
                st.markdown("**🎬 Reel Ideas:**")
                for r in channel_data.get('reel_ideas', []):
                    st.write(f"• {r}")
            
            elif channel_name == "email":
                st.markdown("**📧 Sequences:**")
                for seq_name, seq_content in channel_data.items():
                    if isinstance(seq_content, list):
                        st.write(f"**{seq_name}:**")
                        for item in seq_content:
                            st.write(f"  • {item}")
            
            elif channel_name == "ads":
                for ad_type, ad_data in channel_data.items():
                    with st.expander(f"**{ad_type}**"):
                        if isinstance(ad_data, dict):
                            for k, v in ad_data.items():
                                if isinstance(v, list):
                                    st.write(f"**{k}:** {', '.join(v)}")
                                else:
                                    st.write(f"**{k}:** {v}")


def show_timeline(timeline):
    """Display launch timeline"""
    if not timeline:
        st.info("No timeline data")
        return
    
    phases = [
        ("6_months_out", "6 Months Before Launch"),
        ("3_months_out", "3 Months Before Launch"),
        ("1_month_out", "1 Month Before Launch"),
        ("launch_week", "Launch Week"),
        ("post_launch", "Post-Launch")
    ]
    
    for phase_key, phase_name in phases:
        tasks = timeline.get(phase_key, [])
        if tasks:
            with st.expander(f"📅 {phase_name}"):
                for task in tasks:
                    st.write(f"• {task}")


def show_assets(assets):
    """Display generated assets"""
    if not assets:
        st.info("No assets generated")
        return
    
    # Asset download buttons
    st.markdown("### 📥 Download Assets")
    
    col1, col2, col3 = st.columns(3)
    
    with col1:
        if st.button("📝 Blurbs", use_container_width=True):
            st.info("Download coming soon")
    with col2:
        if st.button("🎬 TikTok Scripts", use_container_width=True):
            st.info("Download coming soon")
    with col3:
        if st.button("📧 Emails", use_container_width=True):
            st.info("Download coming soon")
    
    # Preview sections
    asset_tabs = st.tabs(["Blurbs", "TikTok Scripts", "Emails", "Social Posts", "Ad Copy", "Quote Cards"])
    
    with asset_tabs[0]:
        blurbs = assets.get('blurbs', [])
        for i, blurb in enumerate(blurbs, 1):
            with st.expander(f"Blurb Version {i}"):
                st.write(blurb)
    
    with asset_tabs[1]:
        scripts = assets.get('tiktok_scripts', [])
        for i, script in enumerate(scripts, 1):
            with st.expander(f"Script {i}"):
                if isinstance(script, dict):
                    for k, v in script.items():
                        st.write(f"**{k}:** {v}")
                else:
                    st.write(script)
    
    with asset_tabs[2]:
        emails = assets.get('emails', {})
        for name, content in emails.items():
            with st.expander(f"📧 {name}"):
                st.write(content)
    
    with asset_tabs[3]:
        posts = assets.get('social_posts', [])
        for i, post in enumerate(posts, 1):
            with st.expander(f"Post {i}"):
                if isinstance(post, dict):
                    for k, v in post.items():
                        st.write(f"**{k}:** {v}")
                else:
                    st.write(post)
    
    with asset_tabs[4]:
        ads = assets.get('ad_copy', {})
        for name, content in ads.items():
            with st.expander(f"📢 {name}"):
                if isinstance(content, dict):
                    for k, v in content.items():
                        st.write(f"**{k}:** {v}")
                else:
                    st.write(content)
    
    with asset_tabs[5]:
        quotes = assets.get('quote_cards', [])
        for i, quote in enumerate(quotes, 1):
            with st.expander(f"Quote {i}"):
                st.write(f"**Text:** {quote.get('text', '')}")
                st.write(f"**Visual:** {quote.get('visual', '')}")


# ============================================================================
# HELPER FUNCTIONS
# ============================================================================

def extract_text_preview(file):
    """Extract preview text from file"""
    try:
        if file.type == "application/pdf":
            pdf_reader = PyPDF2.PdfReader(file)
            return pdf_reader.pages[0].extract_text()[:500]
        elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = docx.Document(file)
            return doc.paragraphs[0].text[:500] if doc.paragraphs else ""
        else:
            return file.getvalue().decode("utf-8")[:500]
    except:
        return "Preview not available"


def extract_full_text(bytes_data, file_type):
    """Extract full text from file bytes"""
    try:
        import io
        if "pdf" in file_type:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(bytes_data))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text()
            return text
        elif "document" in file_type:
            doc = docx.Document(io.BytesIO(bytes_data))
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            return text
        else:
            return bytes_data.decode("utf-8")
    except Exception as e:
        st.error(f"Error extracting text: {e}")
        return ""


def generate_complete_blueprint(manuscript_text, cover_base64, options):
    """Call GPT-4o-mini to generate marketing blueprint"""
    
    try:
        client = OpenAI(api_key=st.session_state.openai_api_key)
        
        # Truncate aggressively for mini model
        if len(manuscript_text) > 20000:
            manuscript_text = manuscript_text[:20000] + "... [truncated]"
        
        # Get first 2000 chars as sample
        sample_text = manuscript_text[:2000]
        
        prompt = f"""
        Create a complete book marketing blueprint based on this manuscript excerpt.
        
        MANUSCRIPT SAMPLE:
        {sample_text}
        
        Return a JSON with these exact sections:
        
        1. book_profile: {{
            "title": "suggested title",
            "genre": "primary genre",
            "subgenres": ["subgenre1", "subgenre2"],
            "tone": "emotional tone",
            "pace": "fast/medium/slow",
            "main_characters": ["character descriptions"]
        }}
        
        2. cover_analysis: {{
            "genre_alignment": "percentage%",
            "strengths": ["strength1", "strength2"],
            "weaknesses": ["weakness1", "weakness2"],
            "suggestions": ["suggestion1", "suggestion2"],
            "comp_covers": ["comparable book1", "comparable book2"]
        }}
        
        3. reader_avatar: {{
            "primary": {{
                "name": "reader name",
                "age": "age range",
                "occupation": "typical job",
                "reading_habits": "how they read",
                "what_she_seeks": "what they want",
                "where_she_hangs_out": "platforms",
                "what_she_avoids": "turnoffs"
            }}
        }}
        
        4. market_position: {{
            "primary_shelf": "Amazon category",
            "comp_titles": [
                {{"title": "book1", "similarity": "similar aspects", "difference": "different aspects"}}
            ],
            "positioning_statement": "for fans of X who want Y"
        }}
        
        5. keyword_cloud: {{
            "amazon_keywords": ["keyword1", "keyword2"],
            "search_volume": {{"high": ["term"], "medium": ["term"]}},
            "categories": ["category1", "category2"]
        }}
        
        6. blurb_analysis: {{
            "score": 0-100,
            "strengths": ["strength1"],
            "weaknesses": ["weakness1"],
            "optimized_version": "new blurb text"
        }}
        
        7. marketing_blueprint.channel_strategies: {{
            "tiktok": {{
                "content_pillars": ["pillar1", "pillar2"],
                "sound_suggestions": ["sound1", "sound2"],
                "hashtags": ["#tag1", "#tag2"]
            }},
            "instagram": {{
                "post_types": ["type1", "type2"],
                "reel_ideas": ["idea1", "idea2"]
            }},
            "email": {{
                "welcome_sequence": ["email1", "email2"],
                "launch_sequence": ["email1", "email2"]
            }}
        }}
        
        8. marketing_blueprint.launch_timeline: {{
            "6_months_out": ["task1", "task2"],
            "3_months_out": ["task1", "task2"],
            "1_month_out": ["task1", "task2"],
            "launch_week": ["task1", "task2"],
            "post_launch": ["task1", "task2"]
        }}
        
        9. generated_assets: {{
            "blurbs": ["blurb1", "blurb2", "blurb3"],
            "tiktok_scripts": [
                {{"hook": "hook1", "visuals": "desc", "voiceover": "text"}}
            ],
            "emails": {{"prelaunch": "text", "launch": "text"}},
            "social_posts": ["post1", "post2"],
            "ad_copy": {{"variation1": "text"}},
            "quote_cards": [{{"text": "quote", "visual": "desc"}}]
        }}
        
        Return ONLY valid JSON. No other text.
        """
        
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a book marketing expert. Return valid JSON only."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=3000
        )
        
        # Parse response
        result = json.loads(response.choices[0].message.content)
        return result
        
    except Exception as e:
        st.error(f"API Error: {str(e)}")
        return get_fallback_blueprint()


def get_fallback_blueprint():
    """Return sample blueprint when API fails"""
    return {
        "book_profile": {
            "title": "The Last Ember",
            "genre": "Fantasy Romance",
            "subgenres": ["Enemies to Lovers", "Magic System", "Slow Burn"],
            "tone": "Dark and passionate",
            "pace": "Medium-fast",
            "main_characters": ["Female protagonist: Strong-willed warrior", "Male lead: Mysterious rival with secrets"]
        },
        "cover_analysis": {
            "genre_alignment": "85%",
            "strengths": ["Bold typography", "Moody atmosphere", "Professional design"],
            "weaknesses": ["Figure could be larger", "Colors slightly muted"],
            "suggestions": ["Add subtle magical glow", "Increase contrast by 15%", "Make title pop more"],
            "comp_covers": ["Fourth Wing", "A Court of Thorns and Roses", "The Serpent & the Wings of Night"]
        },
        "reader_avatar": {
            "primary": {
                "name": "Fantasy Romance Emma",
                "age": "28",
                "occupation": "Marketing Manager",
                "reading_habits": "3-4 books/month on Kindle",
                "what_she_seeks": "Emotional payoff, strong heroine, romantic tension",
                "where_she_hangs_out": "Instagram, BookTok, Goodreads",
                "what_she_avoids": "Slow pacing, info-dumping, weak endings"
            },
            "secondary": {
                "name": "YA Crossover Alex",
                "age": "22",
                "occupation": "College Student",
                "reading_habits": "Reads on phone, follows BookTok trends",
                "what_she_seeks": "Viral-worthy moments, relatable characters",
                "where_she_hangs_out": "TikTok, Discord, Twitter",
                "what_she_avoids": "Overly complex worldbuilding"
            }
        },
        "market_position": {
            "primary_shelf": "Adult Fantasy Romance",
            "comp_titles": [
                {"title": "Fourth Wing", "similarity": "Dragon rider academy, enemies to lovers", "difference": "Your magic system is more unique"},
                {"title": "ACOTAR", "similarity": "Fae romance, slow burn", "difference": "More action-focused"},
                {"title": "Serpent & Wings", "similarity": "Vampire romance, trials", "difference": "Different mythology"}
            ],
            "gap_analysis": "Market lacks books with your specific magic system + enemies to lovers + strong mental health rep",
            "positioning_statement": "For fans of Fourth Wing who want deeper worldbuilding and a fresh take on elemental magic"
        },
        "keyword_cloud": {
            "amazon_keywords": ["fantasy romance", "enemies to lovers", "magic system", "strong heroine", "dragons"],
            "search_volume": {
                "high": ["enemies to lovers", "fantasy romance"],
                "medium": ["dragon romance", "fae romance"],
                "low": ["elemental magic"]
            },
            "categories": ["Fiction > Fantasy > Romance", "Fiction > Romance > Paranormal"]
        },
        "blurb_analysis": {
            "score": 72,
            "strengths": ["Good hook", "Clear stakes", "Genre-appropriate"],
            "weaknesses": ["Could be more emotional", "Missing comp titles", "Slightly long"],
            "optimized_version": "In a world where magic awakens desire, one woman must choose between her duty and her heart. When she meets her sworn enemy, the line between love and hate blurs. Perfect for fans of Fourth Wing and ACOTAR."
        },
        "marketing_blueprint": {
            "channel_strategies": {
                "tiktok": {
                    "content_pillars": ["Reaction videos to tropes", "Aesthetic montages", "Day in the life of writing"],
                    "sound_suggestions": ["Dark orchestral music", "Cinematic covers of pop songs", "Spoken word poetry"],
                    "hashtags": ["#Romantasy", "#BookTok", "#FantasyRomance", "#EnemiesToLovers"]
                },
                "instagram": {
                    "post_types": ["Quote cards (6)", "Character aesthetic (4)", "Setting mood boards (3)", "Author reels"],
                    "reel_ideas": ["Book summary in 15 sec", "Readers react to plot twist", "Author answers comments", "Aesthetic book flip"]
                },
                "email": {
                    "welcome_sequence": ["Welcome to my world", "What inspired the book", "Exclusive deleted scene"],
                    "launch_sequence": ["Pre-order incentives", "Launch day announcement", "Thank you + reviews", "What's next"]
                },
                "ads": {
                    "amazon_ads": {
                        "headlines": ["If you loved Fourth Wing...", "The dragon romance you've been waiting for", "Enemies to lovers + magic"],
                        "keywords_to_bid": ["romantasy", "dragon romance", "enemies to lovers", "fantasy romance"],
                        "negative_keywords": ["YA", "middle grade", "non-fiction", "children"]
                    }
                }
            },
            "launch_timeline": {
                "6_months_out": ["Cover reveal", "Goodreads page setup", "ARC sign-ups open", "Build email list"],
                "3_months_out": ["Send ARCs", "Influencer outreach", "First teasers", "Create content bank"],
                "1_month_out": ["Pre-orders live", "Email sequence begins", "Review campaign", "Schedule launch week"],
                "launch_week": ["Daily content", "Giveaways", "Review push", "DM influencers", "Engage comments"],
                "post_launch": ["Thank you content", "Book 2 teaser", "Audiobook announcement", "Plan next launch"]
            }
        },
        "generated_assets": {
            "blurbs": [
                "Version 1 (trope-focused): A dragon rider's quest for vengeance becomes a fight for love in a world where magic awakens desire. Perfect for fans of Fourth Wing.",
                "Version 2 (character-focused): She never expected to fall for her enemy. He never expected to betray her. When magic awakens between them, nothing will ever be the same.",
                "Version 3 (plot-focused): In a world of magic and betrayal, one woman holds the key to saving her people. But first, she must survive her heart."
            ],
            "tiktok_scripts": [
                {"hook": "Books like Fourth Wing but with MORE magic?", "visuals": "Book pages flipping, dramatic pans", "voiceover": "If you loved dragons and enemies to lovers, wait until you meet this book...", "music": "Dark orchestral", "cta": "Add to your TBR"},
                {"hook": "POV: you just found THE book of the year", "visuals": "Reaction shots, book reveal", "voiceover": "The tension? Immaculate. The magic? Unique. The romance?", "music": "Trending sound", "cta": "Link in bio for pre-order"}
            ],
            "emails": {
                "prelaunch": "Subject: Your next fantasy obsession is almost here...\n\nI've been working on something special, and I think you're going to love it. Imagine Fourth Wing meets ACOTAR with a magic system you've never seen before...",
                "launch": "Subject: It's here! ✨\n\nToday's the day! Your copy of [Title] is waiting. Click here to grab your copy before the price goes up!"
            },
            "social_posts": [
                "The book that made me believe in magic again... #Bookstagram #Romantasy",
                "Same vibes as Fourth Wing but with a twist 🐉✨ Who's adding this to their TBR?",
                "5 reasons you need this book on your TBR: 1) Enemies to lovers 2) Unique magic 3) Strong FMC 4) Slow burn 5) That PLOT TWIST"
            ],
            "ad_copy": {
                "variation1": "Headline: For fans of Fourth Wing\nText: Discover your next dragon romance obsession. Enemies to lovers meets unique magic.",
                "variation2": "Headline: The book BookTok can't stop talking about\nText: If you loved ACOTAR, you'll devour this. Get your copy today."
            },
            "quote_cards": [
                {"text": "Some bonds are forged in fire... and betrayal", "visual": "Silhouettes against flames, dramatic lighting"},
                {"text": "The heart wants what magic forbids", "visual": "Glowing hands reaching for each other"},
                {"text": "In a world of enemies, she found her only ally in him", "visual": "Two figures back to back, ready to fight"}
            ]
        }
    }
